import gensim
from gensim.models.doc2vec import Doc2Vec, TaggedDocument
import jieba
import os
from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

TaggedDocument = gensim.models.doc2vec.TaggedDocument


def get_corpus():
    sentences=[]
    for root, dirs, files in os.walk('data/raw_resumes'):
        for name in files:
            file_path = os.path.join(root, name)
            text = read_docx(file_path)
            sentences.append(text)

    train_docs = []
    for i, text in enumerate(sentences):
        word_list = text.split(' ')
        length = len(word_list)
        word_list[length - 1] = word_list[length - 1].strip()
        document = TaggedDocument(word_list, tags=[i])
        train_docs.append(document)

    return train_docs


def train(x_train, vector_size=200, epoch_num=1):
    model_dm = Doc2Vec(vector_size=vector_size, min_count=1, window=3, sample=1e-3, negative=5, workers=4)
    model_dm.build_vocab(x_train)
    model_dm.train(x_train, total_examples=model_dm.corpus_count, epochs=70)
    model_dm.save('model_doc2vec')
    return model_dm



if __name__ == '__main__':
    x_train = get_corpus()
    model_dm = train(x_train)
    model_dm = Doc2Vec.load("model_doc2vec")
    with open('JobDescription.txt', 'r') as f:
        text_test = f.read()
    text_cut = jieba.cut(text_test)
    text_raw = []

    for i in list(text_cut):
        text_raw.append(i)

    inferred_vector_dm = model_dm.infer_vector(text_raw)
    sims = model_dm.docvecs.most_similar([inferred_vector_dm], topn=10)

    for count, sim in sims:
        sentence = x_train[count]
        words = ''

        for word in sentence[0]:
            words = words + word + ' '

        print(words, sim, len(sentence[0]))
        print('------------------------------------------------------------------------------')